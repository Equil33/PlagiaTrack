#!/usr/bin/env python3
import sys
import os
import json
import re
import numpy as np
from datetime import datetime
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from difflib import SequenceMatcher
import Levenshtein
import textstat
from rapidfuzz import fuzz
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet
import os
import os

def convert_pdf_to_txt(path):
    try:
        import PyPDF2
        with open(path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            text = ""
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\\n"
            return text
    except Exception as e:
        print(f"Erreur conversion PDF {path}: {e}", file=sys.stderr)
        return ""

def convert_docx_to_txt(path):
    try:
        import docx
        doc = docx.Document(path)
        text = "\\n".join([p.text for p in doc.paragraphs])
        return text
    except Exception as e:
        print(f"Erreur conversion DOCX {path}: {e}", file=sys.stderr)
        return ""

def lire_fichier(path):
    ext = os.path.splitext(path)[1].lower()
    if ext == '.txt':
        try:
            with open(path, 'r', encoding='utf-8-sig') as f:
                return f.read()
        except Exception as e:
            print(f"Erreur lecture fichier {path}: {e}", file=sys.stderr)
            return ""
    elif ext == '.pdf':
        return convert_pdf_to_txt(path)
    elif ext == '.docx':
        return convert_docx_to_txt(path)
    else:
        print(f"Format non supporté pour {path}", file=sys.stderr)
        return ""

def nettoyage_texte(texte):
    texte = texte.lower()
    texte = re.sub(r'[^a-zàâçéèêëîïôûùüÿñæœ\\s]', ' ', texte)
    texte = re.sub(r'\\s+', ' ', texte)
    return texte.strip()

def tfidf_cosine_sim(texts):
    vectorizer = TfidfVectorizer()
    tfidf_matrix = vectorizer.fit_transform(texts)
    sim_matrix = cosine_similarity(tfidf_matrix)
    return sim_matrix

def levenshtein_sim(texts):
    n = len(texts)
    sim_matrix = np.zeros((n, n))
    for i in range(n):
        for j in range(i, n):
            dist = Levenshtein.distance(texts[i], texts[j])
            max_len = max(len(texts[i]), len(texts[j]))
            sim = 1 - dist / max_len if max_len > 0 else 1
            sim_matrix[i][j] = sim
            sim_matrix[j][i] = sim
    return sim_matrix

def jaccard_sim(texts, ngram=3):
    def ngrams(text, n):
        return set([text[i:i+n] for i in range(len(text)-n+1)])
    n = len(texts)
    sim_matrix = np.zeros((n, n))
    for i in range(n):
        ngrams_i = ngrams(texts[i], ngram)
        for j in range(i, n):
            ngrams_j = ngrams(texts[j], ngram)
            inter = len(ngrams_i.intersection(ngrams_j))
            union = len(ngrams_i.union(ngrams_j))
            sim = inter / union if union > 0 else 1
            sim_matrix[i][j] = sim
            sim_matrix[j][i] = sim
    return sim_matrix

def complexite_lexicale(texte):
    mots = texte.split()
    if not mots:
        return 0
    mots_rares = sum(1 for mot in mots if textstat.dale_chall_readability_score(mot) > 7)
    return mots_rares / len(mots)

def longueur_moyenne_phrase(texte):
    phrases = re.split(r'[.!?]+', texte)
    phrases = [p.strip() for p in phrases if p.strip()]
    if not phrases:
        return 0
    mots_par_phrase = [len(p.split()) for p in phrases]
    return sum(mots_par_phrase) / len(phrases)

def repetition_structures(texte):
    mots = texte.split()
    if len(mots) < 2:
        return 0
    repetitions = 0
    for i in range(len(mots)-1):
        if fuzz.ratio(mots[i], mots[i+1]) > 90:
            repetitions += 1
    return repetitions / (len(mots)-1)

def burstiness(texte):
    phrases = re.split(r'[.!?]+', texte)
    phrases = [p.strip() for p in phrases if p.strip()]
    longueurs = [len(p.split()) for p in phrases]
    if not longueurs:
        return 0
    return np.std(longueurs) / np.mean(longueurs)

def indice_suspicion_ia(texte):
    c_lex = complexite_lexicale(texte)
    lmp = longueur_moyenne_phrase(texte)
    rep = repetition_structures(texte)
    burst = burstiness(texte)
    score = (c_lex * 0.3 + (1 - lmp/20) * 0.2 + rep * 0.3 + burst * 0.2) * 100
    return max(0, min(100, score))

def detection_sections_similaires(text1, text2, seuil=0.7):
    matcher = SequenceMatcher(None, text1, text2)
    blocs = []
    for bloc in matcher.get_matching_blocks():
        if bloc.size > 50:
            ratio = matcher.ratio()
            if ratio >= seuil:
                extrait1 = text1[bloc.a:bloc.a+bloc.size]
                extrait2 = text2[bloc.b:bloc.b+bloc.size]
                blocs.append({
                    'extrait_document_1': extrait1,
                    'extrait_document_2': extrait2,
                    'similarite': ratio
                })
    return blocs

class PDFReport:
    def __init__(self, chemin_pdf):
        self.chemin_pdf = chemin_pdf
        self.story = []
        self.styles = getSampleStyleSheet()
        self.doc = SimpleDocTemplate(chemin_pdf, pagesize=letter)

    def header(self, canvas, doc):
        canvas.saveState()
        canvas.setFont('Helvetica-Bold', 12)
        canvas.drawCentredString(letter[0] / 2.0, letter[1] - 40, "Rapport de Detection de Plagiat")
        canvas.restoreState()

    def footer(self, canvas, doc):
        canvas.saveState()
        canvas.setFont('Helvetica-Oblique', 7)
        page_num = f"Page {doc.page}"
        canvas.drawCentredString(letter[0] / 2.0, 20, page_num)
        canvas.restoreState()

    def add_paragraph(self, text, style_name='Normal'):
        paragraph = Paragraph(text, self.styles[style_name])
        self.story.append(paragraph)
        self.story.append(Spacer(1, 12))

    def add_table(self, data, col_widths):
        # Wrap long text in cells using Paragraph with word wrap
        wrapped_data = []
        for row in data:
            wrapped_row = []
            for item in row:
                # Use Paragraph to enable wrapping
                wrapped_item = Paragraph(str(item), self.styles['Normal'])
                wrapped_row.append(wrapped_item)
            wrapped_data.append(wrapped_row)

        table = Table(wrapped_data, colWidths=col_widths)
        style = TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ])
        table.setStyle(style)
        self.story.append(table)
        self.story.append(Spacer(1, 12))

    def build_pdf(self):
        self.doc.build(self.story, onFirstPage=self.header, onLaterPages=self.header)

def generer_pdf(rapport_complet, chemin_pdf):
    pdf = PDFReport(chemin_pdf)

    pdf.add_paragraph('Informations Generales', 'Heading1')
    pdf.add_paragraph(f"Numero de rapport : {rapport_complet['rapport_id']}")
    pdf.add_paragraph(f"Date et heure : {rapport_complet['analyse_effectuee_le']}")
    pdf.add_paragraph(f"Nombre de documents analyses : {len(rapport_complet['documents_compares'])}")
    pdf.add_paragraph(f"Nombre de comparaisons effectuees : {len(rapport_complet['sections_similaires'])}")

    pdf.add_paragraph('Tableau recapitulatif des comparaisons', 'Heading1')
    headers = ['Fichier 1', 'Indice IA Fichier 1 (%)', 'Fichier 2', 'Indice IA Fichier 2 (%)', 'Indice de plagiat (%)', 'Conclusion']
    col_widths = [80, 80, 80, 80, 80, 80]

    table_data = [headers]
    if not rapport_complet['sections_similaires']:
        unique_files = set(rapport_complet.get('documents_compares', []))
        for file in unique_files:
            filename = os.path.basename(file)
            indice_ia_val = round(rapport_complet.get('indice_suspicion_IA', 0), 2)
            table_data.append([filename, f"{indice_ia_val}%", '-', '-', '0.01%', 'Aucun plagiat'])
    else:
        # Trier les sections similaires par score de similarite décroissant
        sections_triees = sorted(rapport_complet['sections_similaires'], key=lambda x: x.get('score_similarite', 0), reverse=True)
        for comparaison in sections_triees:
            fichier1 = os.path.basename(comparaison['document1'])
            fichier2 = os.path.basename(comparaison['document2'])
            indice_plagiat = round(comparaison.get('score_similarite', 0) * 100, 2)
            # Calculate IA index for each file individually
            # Find IA index for fichier1 and fichier2 from documents_compares and indices_ia
            # Since indices_ia is not passed here, we calculate on the fly:
            # We need to map documents_compares to indices_ia
            indices_ia_map = {}
            for idx, doc in enumerate(rapport_complet.get('documents_compares', [])):
                texte = lire_fichier(doc)
                indices_ia_map[os.path.basename(doc)] = round(indice_suspicion_ia(texte), 2)
            ia_fichier1 = indices_ia_map.get(fichier1, 0)
            ia_fichier2 = indices_ia_map.get(fichier2, 0)
            conclusion = 'Probable plagiat' if indice_plagiat >= 70 else 'Aucun plagiat'
            table_data.append([fichier1, f"{ia_fichier1}%", fichier2, f"{ia_fichier2}%", f"{indice_plagiat}%", conclusion])

    pdf.add_table(table_data, col_widths)

    # Utiliser les sections triées pour les détails aussi
    for comparaison in sections_triees:
        pdf.add_paragraph(f"Details pour : {os.path.basename(comparaison['document1'])} vs {os.path.basename(comparaison['document2'])}", 'Heading2')
        pdf.add_paragraph(f"Similarite TF-IDF : {round(comparaison.get('score_similarite', 0) * 100, 2)}%")
        pdf.add_paragraph(f"Indice IA : {round(rapport_complet.get('indice_suspicion_IA', 0), 2)}%")

        for section in comparaison.get('sections_similaires', []):
            pdf.add_paragraph("Extrait document 1 :")
            pdf.add_paragraph(section.get('extrait_document_1', ''))
            pdf.add_paragraph("Extrait document 2 :")
            pdf.add_paragraph(section.get('extrait_document_2', ''))
            pdf.add_paragraph(f"Taux de similarite : {round(section.get('similarite', 0) * 100, 2)}%")

    pdf.add_paragraph("Analyse IA par document", 'Heading1')
    for doc in rapport_complet['documents_compares']:
        texte = lire_fichier(doc)
        pdf.add_paragraph(os.path.basename(doc), 'Heading2')
        pdf.add_paragraph(f"Complexite lexicale : {complexite_lexicale(texte):.2f}")
        pdf.add_paragraph(f"Longueur moyenne des phrases : {longueur_moyenne_phrase(texte):.2f}")
        pdf.add_paragraph(f"Repetition de structures : {repetition_structures(texte):.2f}")
        pdf.add_paragraph(f"Indice global de suspicion IA : {indice_suspicion_ia(texte):.2f}%")

    pdf.add_paragraph("Conclusion generale", 'Heading1')
    conclusion_text = "Analyse automatique generee. "
    if any(round(comp.get('score_similarite', 0) * 100, 2) >= 70 for comp in rapport_complet['sections_similaires']):
        conclusion_text += "Des similarites elevees ont ete detectees, indiquant un probable plagiat. "
        conclusion_text += "Une relecture manuelle est recommandee. "
    else:
        conclusion_text += "Aucun plagiat significatif detecte. "
    conclusion_text += "Veuillez verifier les sources et les documents suspects."
    pdf.add_paragraph(conclusion_text)

    pdf.build_pdf()

def main():
    import argparse
    parser = argparse.ArgumentParser(description="Analyse de plagiat")
    parser.add_argument('fichiers', nargs='+', help='Fichiers a analyser')
    parser.add_argument('--seuil', type=float, default=0.7, help='Seuil de similarite (0-1)')
    parser.add_argument('--algo', type=str, default='tfidf', choices=['tfidf', 'levenshtein', 'jaccard', 'lcs'], help='Algorithme de similarite')
    parser.add_argument('--mode', type=str, default='cible', choices=['cible', 'global'], help='Mode d\'analyse')
    args = parser.parse_args()

    fichiers = args.fichiers
    seuil = args.seuil
    algo = args.algo
    mode = args.mode

    textes_bruts = [lire_fichier(f) for f in fichiers]
    textes_nettoyes = [nettoyage_texte(t) for t in textes_bruts]

    if algo == 'tfidf':
        sim_matrix = tfidf_cosine_sim(textes_nettoyes)
    elif algo == 'levenshtein':
        sim_matrix = levenshtein_sim(textes_nettoyes)
    elif algo == 'jaccard':
        sim_matrix = jaccard_sim(textes_nettoyes)
    elif algo == 'lcs':
        sim_matrix = lcs_sim(textes_nettoyes)
    else:
        print(f"Algorithme inconnu: {algo}", file=sys.stderr)
        sys.exit(1)

    indices_ia = [indice_suspicion_ia(t) for t in textes_bruts]

    n = len(fichiers)
    rapports = []
    taux_sim_global = 0
    count_sim = 0

    for i in range(n):
        for j in range(i+1, n):
            score = sim_matrix[i][j]
            if score >= seuil:
                blocs = detection_sections_similaires(textes_bruts[i], textes_bruts[j], seuil)
                rapport_item = {
                    "document1": fichiers[i],
                    "document2": fichiers[j],
                    "score_similarite": float(score),
                    "sections_similaires": blocs
                }
                taux_sim_global += score
                count_sim += 1
            else:
                # For pairs below threshold, add with minimal score and empty sections
                rapport_item = {
                    "document1": fichiers[i],
                    "document2": fichiers[j],
                    "score_similarite": 0.0001,
                    "sections_similaires": []
                }
            rapports.append(rapport_item)

    taux_sim_global = (taux_sim_global / count_sim) * 100 if count_sim > 0 else 0

    rapport_complet = {
        "rapport_id": f"RPT-{datetime.now().strftime('%Y%m%d-%H%M%S')}",
        "documents_compares": fichiers,
        "taux_similarite": round(taux_sim_global, 2),
        "indice_suspicion_IA": round(sum(indices_ia)/len(indices_ia), 2) if indices_ia else 0,
        "sections_similaires": rapports,
        "analyse_effectuee_le": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        "statut": "Plagiat probable" if taux_sim_global >= seuil else "Pas de plagiat detecte",
        "recommendation": "Verification manuelle recommandee." if taux_sim_global >= seuil else "Aucun probleme detecte."
    }

    reports_dir = os.path.join(os.path.dirname(__file__), '..', 'reports')
    if not os.path.exists(reports_dir):
        os.makedirs(reports_dir)

    chemin_json = f"rapport_session_{datetime.now().strftime('%Y-%m-%d_%H-%M-%S')}.json"
    chemin_json_full = os.path.join(reports_dir, chemin_json)
    with open(chemin_json_full, 'w', encoding='utf-8') as f:
        json.dump(rapport_complet, f, ensure_ascii=False, indent=4)

    chemin_pdf = chemin_json.replace('.json', '.pdf')
    chemin_pdf_full = os.path.join(reports_dir, chemin_pdf)
    generer_pdf(rapport_complet, chemin_pdf_full)

    print(json.dumps(rapport_complet, ensure_ascii=False, indent=4))

if __name__ == "__main__":
    main()
